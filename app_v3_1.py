import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path

import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates"

TEMPLATES = {
    "Yaounde – 8:50 Uhr": {
        "file": "yaounde_0850.docx",
        "ort": "Yaounde",
        "niveau": "Mündliche Prüfung B2",
        "pruefungsnummer": "9217259",
        "datum": "01.06.2026",
        "farbe": "ROT",
    },
    "Yaounde – 11:10 Uhr": {
        "file": "yaounde_1110.docx",
        "ort": "Yaounde",
        "niveau": "Mündliche Prüfung B2",
        "pruefungsnummer": "9217783",
        "datum": "01.05.2026",
        "farbe": "GRÜN",
    },
    "Douala – 8:50 Uhr": {
        "file": "douala_0850.docx",
        "ort": "Douala",
        "niveau": "Mündliche Prüfung B2",
        "pruefungsnummer": "9217250",
        "datum": "29.05.2026",
        "farbe": "ROT",
    },
    "Douala – 11:10 Uhr": {
        "file": "douala_1110.docx",
        "ort": "Douala",
        "niveau": "Mündliche Prüfung B2",
        "pruefungsnummer": "9217256",
        "datum": "28.05.2026",
        "farbe": "GRÜN",
    },
}

st.set_page_config(
    page_title="Mündliche Prüfung – TN-Liste",
    page_icon="📄",
    layout="centered",
)

st.title("Mündliche Prüfung – TN-Liste")
st.caption("Vorlage auswählen, Prüfungsdaten anpassen, TN-Liste einfügen und Word-Datei erstellen.")


def reset_defaults():
    name = st.session_state["template_choice"]
    data = TEMPLATES[name]
    st.session_state["ort"] = data["ort"]
    st.session_state["niveau"] = data["niveau"]
    st.session_state["pruefungsnummer"] = data["pruefungsnummer"]
    st.session_state["datum"] = data["datum"]
    st.session_state["farbe"] = data["farbe"]


if "template_choice" not in st.session_state:
    st.session_state["template_choice"] = "Yaounde – 11:10 Uhr"
    reset_defaults()

st.subheader("1. Vorlage auswählen")
st.selectbox(
    "Vorlage",
    list(TEMPLATES.keys()),
    key="template_choice",
    on_change=reset_defaults,
)

st.subheader("2. Prüfungsdaten")
col1, col2 = st.columns(2)

with col1:
    st.text_input("Ort", key="ort")
    st.text_input("Niveau", key="niveau")
    st.text_input("Prüfungsnummer", key="pruefungsnummer")

with col2:
    st.text_input("Datum", key="datum", placeholder="TT.MM.JJJJ")
    st.selectbox(
        "Farbe",
        ["GRÜN", "ROT", "GELB", "BLAU"],
        key="farbe",
    )

st.subheader("3. TN-Liste einfügen")
tn_text = st.text_area(
    "TN-Liste",
    height=330,
    placeholder=(
        "GK\nGinette Kouowou\nD384291\nInitial\n"
        "PM\nPamela Marthe Huguette MBACK\nD384295\nInitial"
    ),
)


def parse_tn_list(raw_text):
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    participants = []
    i = 0

    while i < len(lines):
        if (
            i + 2 < len(lines)
            and re.fullmatch(r"[A-Za-zÄÖÜäöüßÀ-ÿ]{1,6}", lines[i])
            and re.fullmatch(r"D\d+", lines[i + 2], flags=re.IGNORECASE)
        ):
            participants.append(lines[i + 1])
            i += 3

            # Optionaler Status, z.B. Initial / In Lobby
            if i < len(lines) and not re.fullmatch(
                r"[A-Za-zÄÖÜäöüßÀ-ÿ]{1,6}", lines[i]
            ):
                i += 1
            elif i < len(lines) and lines[i].lower() in {
                "initial", "in lobby", "lobby", "ready"
            }:
                i += 1
        else:
            i += 1

    return participants


def make_pairs(participants):
    pairs = []
    for i in range(0, len(participants), 2):
        first = participants[i]
        second = participants[i + 1] if i + 1 < len(participants) else None
        if second:
            pairs.append((f"{first} (1)", f"{second} (2)"))
        else:
            pairs.append((f"{first} (1)", None))
    return pairs


def set_run_font(run, size, bold=None):
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph.clear()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)
    return run


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def write_cell(cell, lines, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    # Keep the existing cell/table geometry, but normalize text formatting.
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, size, bold)
        if idx < len(lines) - 1:
            run.add_break()


def format_header(doc, ort, niveau, pruefungsnummer, datum, farbe):
    # Vor den sichtbaren Kopfzeilen können leere Word-Absätze stehen.
    visible_paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    if len(visible_paragraphs) < 4:
        raise ValueError("Die vier Kopfzeilen der Vorlage konnten nicht erkannt werden.")

    values = [ort, niveau, pruefungsnummer, f"{datum} – {farbe}"]
    for paragraph, value in zip(visible_paragraphs[:4], values):
        set_paragraph_text(
            paragraph, value, size=14, bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER
        )


def build_document(template_path, participants, ort, niveau, pruefungsnummer, datum, farbe):
    doc = Document(template_path)

    if not doc.tables:
        raise ValueError("In der Vorlage wurde keine Tabelle gefunden.")

    format_header(doc, ort, niveau, pruefungsnummer, datum, farbe)

    table = doc.tables[0]
    pairs = make_pairs(participants)

    # Header row: 14 pt bold
    if table.rows:
        for cell in table.rows[0].cells:
            write_cell(
                cell,
                [cell.text.strip()],
                size=14,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

    pair_index = 0

    for row in table.rows[1:]:
        if len(row.cells) < 3:
            continue

        left, middle, right = row.cells[0], row.cells[1], row.cells[2]
        row_text = " ".join(c.text.strip() for c in row.cells)
        is_pause = "PAUSE" in row_text.upper()

        # Times: 14 pt bold
        left_text = left.text.strip()
        right_text = right.text.strip()
        write_cell(left, [left_text], size=14, bold=True)
        write_cell(right, [right_text], size=14, bold=True)

        if is_pause:
            # Pause centered, 12 pt normal. Yellow highlighting only on time cells.
            write_cell(middle, ["PAUSE"], size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(left, "FFF200")
            set_cell_shading(right, "FFF200")
            continue

        # Normal participant rows: names 12 pt, normal, two lines in same cell.
        if pair_index < len(pairs):
            first, second = pairs[pair_index]
            lines = [first] + ([second] if second else [])
            write_cell(middle, lines, size=12, bold=False)
            pair_index += 1
        else:
            write_cell(middle, [], size=12, bold=False)

    if pair_index < len(pairs):
        raise ValueError(
            f"Die Vorlage hat nicht genug Teilnehmer-Zeilen. "
            f"{len(pairs)} Paare erkannt, aber nur {pair_index} konnten eingefügt werden."
        )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output, len(pairs)

st.subheader("4. Word-Datei erstellen")

if st.button("Liste erstellen", type="primary", use_container_width=True):
    if not tn_text.strip():
        st.error("Bitte zuerst die TN-Liste einfügen.")
    else:
        participants = parse_tn_list(tn_text)

        if not participants:
            st.error("Es konnten keine Teilnehmer erkannt werden. Bitte TN-Listenformat prüfen.")
        else:
            selected = TEMPLATES[st.session_state["template_choice"]]
            template_path = TEMPLATE_DIR / selected["file"]

            try:
                result, pair_count = build_document(
                    template_path=template_path,
                    participants=participants,
                    ort=st.session_state["ort"].strip(),
                    niveau=st.session_state["niveau"].strip(),
                    pruefungsnummer=st.session_state["pruefungsnummer"].strip(),
                    datum=st.session_state["datum"].strip(),
                    farbe=st.session_state["farbe"],
                )

                filename = (
                    f"{st.session_state['datum'].replace('.', '-')}_"
                    f"{st.session_state['ort']}_"
                    f"{st.session_state['farbe']}_"
                    f"{st.session_state['pruefungsnummer']}.docx"
                )

                st.success(
                    f"Fertig: {len(participants)} Teilnehmer / {pair_count} Paar(e)."
                )

                if len(participants) % 2:
                    st.info("Ungerade Teilnehmerzahl: Der letzte Teilnehmer wurde allein als (1) eingetragen.")

                st.download_button(
                    "Word-Datei herunterladen",
                    data=result,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as exc:
                st.error(str(exc))